#!/usr/bin/env python3
"""Generate a Word (.docx) version of the technical methodology markdown."""

from __future__ import annotations

import re
from pathlib import Path


def _strip_md_bold(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _latex_to_plain(text: str) -> str:
    """Convert lightweight LaTeX in markdown to readable plain text for Word."""
    s = text.strip()
    s = s.replace(r"\!", "")
    s = s.replace(r"\left(", "(").replace(r"\right)", ")")
    s = s.replace(r"\log", "log")
    s = s.replace(r"\max", "max")
    s = s.replace(r"\min", "min")
    s = s.replace(r"\exp", "exp")
    s = s.replace(r"\sum", "Σ")
    s = s.replace(r"\cdot", "·")
    s = s.replace(r"\Delta", "Δ")
    s = s.replace(r"\gamma", "γ")
    s = s.replace(r"\lambda", "λ")
    s = s.replace(r"\mu", "μ")
    s = s.replace(r"\sigma", "σ")
    s = s.replace(r"\beta", "β")
    s = s.replace(r"\eta", "η")
    s = s.replace(r"\widehat{\mathrm{MMR}}", "MMR_hat")
    s = s.replace(r"\mathrm{MMR}", "MMR")
    s = s.replace(r"\mathrm{clamp}", "clamp")
    s = s.replace(r"\log\!", "log")
    s = re.sub(r"\\,\s*", " ", s)
    s = re.sub(r"\\qquad", "    ", s)
    s = re.sub(r"\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "docs" / "Technical_Methodology.md"
    docx_path = root / "docs" / "Technical_Methodology.docx"

    if not md_path.exists():
        raise FileNotFoundError(md_path)

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        ncols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=ncols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(table_rows):
            for c_idx in range(ncols):
                cell_text = row[c_idx] if c_idx < len(row) else ""
                table.rows[r_idx].cells[c_idx].text = _strip_md_bold(cell_text.strip())
        doc.add_paragraph("")
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line.rstrip())
            p.style = "Intense Quote"
            i += 1
            continue

        if line.strip() == "---":
            flush_table()
            doc.add_paragraph("")
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            if re.match(r"^\|\s*-+", line.strip()):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()

        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(_strip_md_bold(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_strip_md_bold(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_strip_md_bold(stripped[2:]), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(_strip_md_bold(stripped[2:]), style="List Bullet")
        elif stripped.startswith("\\[") or (
            stripped.startswith("\\(") and stripped.endswith("\\)")
        ):
            eq_lines = [stripped]
            if stripped.startswith("\\[") and not stripped.endswith("\\]"):
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("\\]"):
                    eq_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    eq_lines.append(lines[i].strip())
            eq_raw = " ".join(eq_lines)
            eq_raw = eq_raw.strip("\\[]()")
            eq_text = _latex_to_plain(eq_raw)
            p = doc.add_paragraph(eq_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.name = "Cambria Math"
        else:
            doc.add_paragraph(_strip_md_bold(stripped))

        i += 1

    flush_table()
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
